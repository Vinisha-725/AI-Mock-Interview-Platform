from typing import List
from io import BytesIO
from models import Project, Experience, Education, ResumeUploadResponse
import re

try:
    import docx
except ImportError:
    docx = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    try:
        from pypdf import PdfReader
    except ImportError:
        PdfReader = None

class ResumeParser:
    def normalize_text(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def contains_keyword(self, text: str, keyword: str) -> bool:
        start_boundary = r'(?<![A-Za-z0-9+#]|(?<=[A-Za-z0-9])\.)'
        end_boundary = r'(?![A-Za-z0-9+#]|\.(?=[A-Za-z0-9]))'
        pattern = start_boundary + re.escape(keyword) + end_boundary
        return re.search(pattern, text, re.IGNORECASE) is not None

    def dedupe(self, values: List[str]) -> List[str]:
        seen = set()
        result = []
        for value in values:
            key = value.lower()
            if key not in seen:
                seen.add(key)
                result.append(value)
        return result

    def get_section(self, text: str, headings: List[str]) -> str:
        heading_pattern = "|".join(re.escape(heading) for heading in headings)
        stop_headings = [
            "summary", "objective", "experience", "work experience", "professional experience",
            "employment", "projects", "education", "certifications", "achievements",
            "awards", "languages", "interests", "contact", "profile"
        ]
        stop_pattern = "|".join(re.escape(heading) for heading in stop_headings if heading not in headings)
        pattern = rf"(?:^|\n)\s*(?:{heading_pattern})\s*:?\s*\n?(.*?)(?=\n\s*(?:{stop_pattern})\s*:?\s*\n|$)"
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        return match.group(1).strip() if match else ""

    def extract_skills_from_section(self, text: str) -> List[str]:
        section = self.get_section(text, ["skills", "technical skills", "technologies", "tech stack"])
        if not section:
            return []

        tokens = re.split(r"[,;|•●▪\-–—\n/]+", section)
        ignored = {
            "skills", "technical skills", "technologies", "tech stack", "tools",
            "languages", "frameworks", "databases", "cloud", "backend", "frontend"
        }
        skills = []
        for token in tokens:
            skill = re.sub(r"\s+", " ", token).strip(" .:()[]{}")
            if ":" in skill:
                skill = skill.split(":")[-1].strip()
            if not skill or skill.lower() in ignored:
                continue
            if len(skill) < 2 or len(skill) > 35:
                continue
            if not re.search(r"[A-Za-z+#.]", skill):
                continue
            skills.append(skill)

        return self.dedupe(skills)[:25]

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        try:
            if PdfReader is None:
                print("Error parsing PDF: install PyPDF2 or pypdf")
                return ""

            pdf_reader = PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            print(f"Extracted {len(text)} characters from PDF")
            return text
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            return ""

    def extract_text_from_docx(self, file_content: bytes) -> str:
        try:
            if docx is None:
                print("Error parsing DOCX: install python-docx")
                return ""

            doc = docx.Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            print(f"Extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return ""

    def extract_text(self, file_content: bytes, filename: str) -> str:
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            text = self.extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            text = self.extract_text_from_docx(file_content)
        else:
            try:
                text = file_content.decode('utf-8')
                print(f"Extracted {len(text)} characters from text file")
            except UnicodeDecodeError:
                text = ""
                print("Could not decode file as text")

        return self.normalize_text(text)

    def extract_skills(self, text: str) -> List[str]:
        # Comprehensive list of technical skills
        skill_keywords = [
            # Programming Languages
            "JavaScript", "Python", "Java", "C++", "C#", "C", "Go", "Rust", "Swift", "Kotlin", "PHP", "Ruby", "Scala", "TypeScript", "R", "MATLAB", "Perl", "Lua", "Haskell",
            # Frontend
            "React", "React.js", "Angular", "Vue", "Vue.js", "Next.js", "Nuxt.js", "Svelte", "Ember", "Backbone", "jQuery", "HTML", "HTML5", "CSS", "CSS3", "SASS", "SCSS", "LESS", "Tailwind", "Tailwind CSS", "Bootstrap", "Vite", "Redux",
            # Backend
            "Node.js", "Node", "Express", "Express.js", "Django", "Flask", "FastAPI", "Spring", "Spring Boot", "Ruby on Rails", "Laravel", "ASP.NET", "NestJS", "Hapi",
            # Databases
            "SQL", "MySQL", "PostgreSQL", "MongoDB", "Mongoose", "Redis", "Elasticsearch", "Cassandra", "DynamoDB", "Firebase", "SQLite", "Oracle", "MariaDB", "Supabase",
            # Cloud & DevOps
            "AWS", "Azure", "GCP", "Google Cloud", "Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins", "GitLab CI", "CircleCI", "Travis CI", "Vercel", "Netlify",
            # Tools
            "Git", "GitHub", "GitLab", "Bitbucket", "Jira", "Confluence", "Slack", "VS Code", "IntelliJ", "Eclipse", "Visual Studio", "Postman", "Figma",
            # APIs & Architecture
            "REST", "RESTful", "GraphQL", "SOAP", "API", "Microservices", "Monolith", "Serverless", "gRPC",
            # Data Science & ML
            "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Jupyter", "Spark", "Hadoop", "OpenAI", "LLM", "LangChain",
            # Mobile
            "React Native", "Flutter", "Ionic", "Xamarin", "Android", "iOS", "SwiftUI", "Jetpack Compose",
            # Other
            "Linux", "Unix", "Bash", "Shell", "PowerShell", "CI/CD", "Agile", "Scrum", "Kanban", "DevOps", "TDD", "BDD",
            "Machine Learning", "Deep Learning", "AI", "Artificial Intelligence", "Data Science", "Big Data", "Blockchain", "MERN", "MERN Stack", "JWT"
        ]

        found_skills = self.extract_skills_from_section(text)

        # Search for skills in entire text with more lenient matching
        text_lower = text.lower()
        for skill in skill_keywords:
            skill_lower = skill.lower()
            # Check if skill appears as a whole word or with common separators
            if skill_lower in text_lower:
                found_skills.append(skill)
            # Also check for partial matches for multi-word skills
            elif ' ' in skill_lower:
                parts = skill_lower.split()
                if all(part in text_lower for part in parts):
                    found_skills.append(skill)

        # Also extract any capitalized words that might be skills (aggressive fallback)
        words = re.findall(r'\b[A-Z][a-zA-Z+#.]*\b', text)
        common_words = {"The", "This", "That", "These", "Those", "Experience", "Education", "Skills", "Projects", "Summary", "Objective", "Profile", "Contact", "Company", "University", "College", "School", "Work", "Professional", "Personal", "Key", "Main", "Major", "Minor", "Bachelor", "Master", "PhD", "Doctor", "Mr", "Mrs", "Ms", "Dr", "Prof", "Inc", "Ltd", "LLC", "Corp", "Co", "Jan", "Feb", "Mar", "Apr", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec", "January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"}
        for word in words:
            if word not in common_words and len(word) >= 2 and len(word) <= 30:
                if word not in found_skills:
                    found_skills.append(word)

        found_skills = self.dedupe(found_skills)
        print(f"Found {len(found_skills)} skills: {found_skills}")
        return found_skills[:50]  # Return up to 50 skills

    def extract_projects(self, text: str) -> List[Project]:
        projects = []
        
        # More flexible project section patterns
        project_patterns = [
            r'project[s]?\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\n\d+\.|\nexperience|\neducation|\nskills|$)',
            r'personal project[s]?\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\n\d+\.|\nexperience|\neducation|\nskills|$)',
            r'key project[s]?\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\n\d+\.|\nexperience|\neducation|\nskills|$)',
        ]
        
        for pattern in project_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                project_text = match.strip()
                if len(project_text) > 30:  # Filter out very short matches
                    # Try to extract tech stack from the project description
                    tech_keywords = ["React", "Node", "Python", "JavaScript", "AWS", "Docker", "SQL", "MongoDB", "Git", "TypeScript", "Java", "Angular", "Vue", "Flask", "Django", "FastAPI"]
                    found_tech = [tech for tech in tech_keywords if self.contains_keyword(project_text, tech)]
                    
                    # Try to extract a project name (first line or before colon)
                    lines = project_text.split('\n')
                    project_name = lines[0].strip() if lines else "Project"
                    if ':' in project_name:
                        project_name = project_name.split(':')[0].strip()
                    
                    projects.append(Project(
                        name=project_name[:50],
                        description=project_text[:300] + "..." if len(project_text) > 300 else project_text,
                        tech=found_tech if found_tech else ["Various Technologies"]
                    ))
        
        print(f"Found {len(projects)} projects")
        return projects[:5]  # Return max 5 projects

    def extract_experience(self, text: str) -> List[Experience]:
        experiences = []
        
        # More flexible experience/work section patterns
        exp_patterns = [
            r'experience\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\neducation|\nprojects|\nskills|$)',
            r'work experience\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\neducation|\nprojects|\nskills|$)',
            r'employment\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\neducation|\nprojects|\nskills|$)',
            r'work history\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\neducation|\nprojects|\nskills|$)',
            r'professional experience\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\neducation|\nprojects|\nskills|$)',
        ]
        
        for pattern in exp_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                exp_text = match.strip()
                if len(exp_text) > 50:
                    # Try to extract company, role, and duration from lines
                    lines = [line.strip() for line in exp_text.split('\n') if line.strip()]
                    
                    for i, line in enumerate(lines):
                        # Look for patterns like "Company Name - Role" or "Role at Company"
                        if len(line) > 15 and i < len(lines) - 1:
                            company = "Company"
                            role = line
                            duration = "Duration"
                            
                            # Try to extract company name
                            if ' at ' in line.lower():
                                parts = line.split(' at ')
                                role = parts[0].strip()
                                company = parts[1].strip()
                            elif ' - ' in line:
                                parts = line.split(' - ')
                                company = parts[0].strip()
                                role = parts[1].strip()
                            elif '|' in line:
                                parts = line.split('|')
                                role = parts[0].strip()
                                company = parts[1].strip() if len(parts) > 1 else "Company"
                            
                            # Try to extract duration from next line
                            if i + 1 < len(lines):
                                next_line = lines[i + 1]
                                if any(char.isdigit() for char in next_line):
                                    duration = next_line
                            
                            experiences.append(Experience(
                                company=company[:50],
                                role=role[:50],
                                duration=duration[:30]
                            ))
                            
                            if len(experiences) >= 5:
                                break
                    if len(experiences) >= 5:
                        break
        
        print(f"Found {len(experiences)} experiences")
        return experiences[:5]  # Return max 5 experiences

    def extract_education(self, text: str) -> List[Education]:
        education_list = []
        
        # More flexible education section patterns
        edu_patterns = [
            r'education\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\nexperience|\nprojects|\nskills|$)',
            r'academic\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\nexperience|\nprojects|\nskills|$)',
            r'qualification\s*:?\s*(.*?)(?=\n\s*\n|\n[A-Z][A-Z]+\s|\nexperience|\nprojects|\nskills|$)',
        ]
        
        for pattern in edu_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                edu_text = match.strip()
                if len(edu_text) > 30:
                    # Try to extract degree, institution, and year
                    lines = [line.strip() for line in edu_text.split('\n') if line.strip()]
                    
                    for line in lines:
                        if len(line) > 20:
                            degree = line
                            institution = "Institution"
                            year = "Year"
                            
                            # Try to extract degree and institution
                            if ' in ' in line.lower() or ' at ' in line.lower():
                                parts = re.split(r'\s+(?:in|at)\s+', line, flags=re.IGNORECASE)
                                if len(parts) >= 2:
                                    degree = parts[0].strip()
                                    institution = parts[1].strip()
                            
                            # Try to extract year
                            year_match = re.search(r'\b(19|20)\d{2}\b', line)
                            if year_match:
                                year = year_match.group()
                            
                            education_list.append(Education(
                                degree=degree[:60],
                                institution=institution[:60],
                                year=year
                            ))
                            
                            if len(education_list) >= 3:
                                break
                    if len(education_list) >= 3:
                        break
        
        print(f"Found {len(education_list)} education entries")
        return education_list[:3]  # Return max 3 education entries

    def extract_tech_stack(self, text: str) -> List[str]:
        # Comprehensive tech stack keywords
        tech_keywords = [
            "JavaScript", "Python", "Java", "C++", "C#", "Go", "Rust", "TypeScript", "PHP", "Ruby", "Swift", "Kotlin",
            "React", "Angular", "Vue", "Next.js", "Nuxt.js", "Svelte", "Ember", "Backbone",
            "Node.js", "Express", "Django", "Flask", "FastAPI", "Spring", "Spring Boot", "Ruby on Rails", "Laravel", "ASP.NET",
            "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch", "Cassandra", "DynamoDB", "Firebase", "SQLite",
            "AWS", "Azure", "GCP", "Google Cloud", "Docker", "Kubernetes", "Terraform", "Ansible",
            "Git", "GitHub", "GitLab", "Bitbucket", "Jenkins", "CI/CD",
            "GraphQL", "REST", "RESTful", "API", "Microservices", "gRPC", "SOAP",
            "Linux", "Unix", "Bash", "Shell", "PowerShell",
            "HTML", "CSS", "SASS", "SCSS", "LESS", "Tailwind", "Bootstrap",
            "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "Pandas", "NumPy", "Spark", "Hadoop"
        ]
        
        found_tech = []
        
        for tech in tech_keywords:
            if self.contains_keyword(text, tech):
                found_tech.append(tech)
        
        print(f"Found {len(found_tech)} tech stack items")
        return found_tech

    def extract_certifications(self, text: str) -> List[str]:
        # More comprehensive certification patterns
        cert_patterns = [
            r'(?:AWS|Azure|GCP|Google Cloud|Microsoft|Oracle|Cisco|Google|Meta|IBM)\s+(?:Certified?|Certificate?|Certification?|Professional|Associate|Architect|Developer|Engineer|Specialist)',
            r'(?:Certified?|Certificate?|Certification?)\s+(?:AWS|Azure|GCP|Google Cloud|Microsoft|Oracle|PMP|Scrum|Agile|CCNA|CCNP|CEH|CISSP|CISA|CISM)',
            r'(?:PMP|Scrum|Agile|CCNA|CCNP|CEH|CISSP|CISA|CISM|ITIL|Prince2)\s+(?:Certified?|Practitioner|Professional)',
            r'(?:Certificate?|Certification?)\s+(?:in|of|for)\s+[^.\n]+',
        ]
        
        found_certs = []
        
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                cert = match.strip()
                if len(cert) > 10 and cert not in found_certs:
                    found_certs.append(cert)
        
        # Also look for lines with certification keywords
        cert_keywords = ["AWS", "Azure", "GCP", "Google Cloud", "Microsoft", "Oracle", "Cisco", "PMP", "Scrum", "Agile", "CCNA", "CCNP", "CEH", "CISSP", "CISA", "CISM", "ITIL"]
        lines = text.split('\n')
        
        for line in lines:
            line_lower = line.lower()
            for cert in cert_keywords:
                if cert.lower() in line_lower and len(line.strip()) > 10:
                    cert_line = line.strip()
                    if cert_line not in found_certs:
                        found_certs.append(cert_line)
                    break
        
        print(f"Found {len(found_certs)} certifications")
        return found_certs[:8]  # Return max 8 certifications

    def parse_resume(self, file_content: bytes, filename: str) -> ResumeUploadResponse:
        print(f"=== Starting resume parsing for: {filename} ===")
        text = self.extract_text(file_content, filename)
        print(f"Total text length: {len(text)} characters")
        print(f"Text preview (first 500 chars): {text[:500]}")

        if len(text) < 30:
            raise ValueError("No readable text found in this resume. If this is a scanned PDF/image, upload a text-based PDF or DOCX.")
        
        # Extract all information
        skills = self.extract_skills(text)
        projects = self.extract_projects(text)
        experience = self.extract_experience(text)
        education = self.extract_education(text)
        tech_stack = self.extract_tech_stack(text)
        certifications = self.extract_certifications(text)
        
        print(f"=== Parsing Complete ===")
        print(f"Skills: {len(skills)}, Projects: {len(projects)}, Experience: {len(experience)}, Education: {len(education)}, Tech Stack: {len(tech_stack)}, Certifications: {len(certifications)}")
        
        return ResumeUploadResponse(
            skills=skills if skills else [],
            projects=projects if projects else [],
            experience=experience if experience else [],
            education=education if education else [],
            techStack=tech_stack if tech_stack else [],
            certifications=certifications if certifications else []
        )

resume_parser = ResumeParser()
